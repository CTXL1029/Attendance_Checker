import os, fitz, copy, subprocess
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# === Bổ sung Endpoint cho Cron-Job / Health Check ===
@app.get("/")
@app.get("/ping")
def health_check():
    return {"status": "active", "message": "Server is alive"}

class AbsentStudent(BaseModel):
    name: str
    permission: bool

class AttendanceRequest(BaseModel):
    class_name: str
    date_str: str       # dd/mm
    date_full: str      # dd-mm-yyyy
    absent_students: list[AbsentStudent]

@app.post("/api/process-attendance")
def process_attendance(data: AttendanceRequest):
    try:
        template_path = os.path.join(os.path.dirname(__file__), "Sample.docx")
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail="Thiếu tệp mẫu Sample.docx")

        doc = Document(template_path)
        num_absent = len(data.absent_students)

        # 1. Giữ nguyên khổ giấy A5 từ Sample.docx & đặt lề mỏng gọn (20pt ~ 0.7cm)
        section = doc.sections[0]
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(25)
        section.right_margin = Pt(25)

        # Đọc chiều cao trang A5 thực tế từ file mẫu (khoảng 420pt nếu A5 Ngang)
        page_height_pt = section.page_height.pt if section.page_height else 420.0
        usable_height = page_height_pt - 40.0  # Trừ 20pt lề trên và 20pt lề dưới

        # 2. Điều chỉnh độ dày hàng linh hoạt theo số học sinh vắng
        if num_absent <= 5:
            row_padding_pt = 6
        elif num_absent <= 10:
            row_padding_pt = 4
        else:
            row_padding_pt = 2

        # 3. Tính toán chính xác khoảng lề trên để toàn bộ khối nằm CHÍNH GIỮA trang A5
        row_height = 14 + (2 * row_padding_pt)
        total_content_height = 65 + 25 + (num_absent * row_height)
        top_space_pt = max(10, int((usable_height - total_content_height) / 2))

        # Áp dụng căn giữa tiêu đề
        for p in doc.paragraphs:
            if "Danh Sách Các Bạn Vắng Mặt" in p.text:
                p.paragraph_format.space_before = Pt(top_space_pt)
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Mulish"

            elif "Ngày dd/mm" in p.text:
                p.text = f"Ngày {data.date_str}"
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Mulish"

        # 4. Xử lý Bảng điểm danh
        if doc.tables:
            table = doc.tables[0]
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            existing_data_rows = len(table.rows) - 1

            # Thêm hoặc xóa hàng cho vừa số học sinh vắng
            if num_absent > existing_data_rows:
                template_tr = table.rows[1]._tr
                for _ in range(num_absent - existing_data_rows):
                    new_tr = copy.deepcopy(template_tr)
                    table._tbl.append(new_tr)
            elif num_absent < existing_data_rows:
                while len(table.rows) - 1 > num_absent:
                    table._tbl.remove(table.rows[-1]._tr)

            # Định dạng hàng tiêu đề bảng
            for cell in table.rows[0].cells:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.name = "Mulish"

            # Điền dữ liệu học sinh vắng
            for index, student in enumerate(data.absent_students, start=1):
                row = table.rows[index]
                cell_values = [
                    str(index),
                    student.name,
                    "Nghỉ có phép" if student.permission else ""
                ]

                for col_idx, text_val in enumerate(cell_values):
                    cell = row.cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    p.paragraph_format.space_before = Pt(row_padding_pt)
                    p.paragraph_format.space_after = Pt(row_padding_pt)
                    
                    run = p.add_run(text_val)
                    run.font.name = "Mulish"
                    run.font.size = Pt(11)

        # Lưu tệp docx tạm thời
        output_docx = f"/tmp/output_{data.date_full}.docx"
        output_pdf_dir = "/tmp"
        output_pdf = f"/tmp/output_{data.date_full}.pdf"
        output_png = f"/tmp/Attendance_Checker_{data.class_name}_{data.date_full}.png"

        doc.save(output_docx)

        # 5. Chuyển đổi sang PDF bằng LibreOffice
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            output_docx, "--outdir", output_pdf_dir
        ], check=True)

        # 6. Chuyển đổi PDF sang PNG bằng PyMuPDF
        pdf_doc = fitz.open(output_pdf)
        page = pdf_doc[0]
        pix = page.get_pixmap(dpi=600)
        pix.save(output_png)
        pdf_doc.close()

        return FileResponse(output_png, media_type="image/png", filename=os.path.basename(output_png))

    except Exception as e:
        print("Lỗi Backend:", str(e))
        raise HTTPException(status_code=500, detail=str(e))